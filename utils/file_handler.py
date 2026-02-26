"""
File type detection and DOCX text extraction.
"""
import io


def detect_file_type(filename: str) -> str:
    """
    Returns: 'pdf', 'image', or 'docx'
    """
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return "pdf"
    elif ext in ("jpg", "jpeg", "png", "webp", "bmp", "tiff"):
        return "image"
    elif ext in ("docx", "doc"):
        return "docx"
    return "unknown"


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract plain text from a DOCX file."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        return "\n\n".join(paragraphs)

    except ImportError:
        return "[python-docx not installed. Install with: pip install python-docx]"
    except Exception as e:
        return f"[DOCX extraction error: {e}]"
